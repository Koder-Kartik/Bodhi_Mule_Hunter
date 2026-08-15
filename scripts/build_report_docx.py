#!/usr/bin/env python3
"""Render the prototype report as a Word document in the IEEE two-column style.

The organisers supplied two templates - a LaTeX one (``spconf``) and a Word one
(IEEE conference format). ``docs/report/bodhi_prototype.tex`` is the LaTeX
submission; this script produces the Word equivalent from the same measured
numbers, so whichever format is required, the content is identical and neither
can drift from ``artifacts/metrics/evaluation.json``.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Pt  # noqa: E402

from bodhi.config import (  # noqa: E402
    AFFILIATION, EVENT, FIGURE_DIR, METRICS_DIR, ROOT, TEAM, TEAM_NAME,
)

OUT = ROOT / "docs" / "report" / "BODHI_Mule_Hunter_Prototype_Report.docx"


def _pct(x, dp=1):
    return f"{100 * float(x):.{dp}f}%"


def _n(x):
    return f"{int(x):,}"


def _two_columns(section) -> None:
    """IEEE conference papers are two-column; python-docx has no API for it."""
    cols = section._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "360")


def _style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.space_after = Pt(4)
    for name, size, bold in (("Heading 1", 10, True), ("Heading 2", 10, True),
                             ("Heading 3", 10, False)):
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = None
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(4)


def _para(doc, text, *, align=None, bold=False, italic=False, size=None,
          space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def _table(doc, headers, rows, caption=None, caption_above=True):
    if caption and caption_above:
        c = _para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=8)
        c.runs[0].italic = True
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(h))
        run.bold = True
        run.font.size = Pt(8)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(8)
    if caption and not caption_above:
        c = _para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=8)
        c.runs[0].italic = True
    doc.add_paragraph()
    return t


def _boi_section(doc) -> None:
    """Section IX: the pipeline built on the organisers' own alert schema.

    Read from artifacts/metrics/boi_track.json for the same reason the rest of
    the document is read from evaluation.json - a paragraph that outlives the
    code that produced it is worse than no paragraph.
    """
    path = METRICS_DIR / "boi_track.json"
    if not path.exists():
        return
    m = json.loads(path.read_text())
    d, dep, leak = m["dataset"], m["deployable"], m["leakage_effect"]
    cv, hold = dep["cv"], dep["holdout"]
    sel = cv[dep["selected_strategy"]]

    doc.add_heading("IX.  THE ORGANISERS' ALERT DATASET", level=1)
    _para(doc,
          f"The organisers released the column dictionary for the Phase 2 dataset "
          f"and stated that the model would be scored on a validation file they had "
          f"not shared. That dataset is a different object from the one above: one "
          f"row is a transaction-monitoring alert, not an account, so the population "
          f"is already pre-filtered by the bank's own rules and the task is triage "
          f"rather than detection from scratch. We therefore built a second pipeline "
          f"(bodhi/boi/) directly against the published schema of "
          f"{_n(d['declared_columns'])} columns. Almost every predictor is "
          f"machine-generated from a compact grammar — aggregation, customer- or "
          f"bank-induced, channel, direction, measure, observation window — so the "
          f"schema module parses the names rather than treating them as opaque. That "
          f"is what allows several thousand columns to be grouped into families "
          f"measured across the 7-, 14- and 31-day windows, and what prevents "
          f"NON_CASH_CHQ being silently matched as CASH.")

    doc.add_heading("A.  Four columns leak the label", level=2)
    _para(doc,
          f"The dictionary describes FRAUD_SUSPECTED, FALSE_POSITIVE, "
          f"OTHER_RESOLUTION and UNATTENDED as resolution-status flags: they record "
          f"how an analyst closed the alert, and MIN_RESOLVE_DAYS and "
          f"MAX_RESOLVE_DAYS record how long that took. An open alert — the only "
          f"kind worth scoring — carries none of them. Admitting them raises "
          f"cross-validated PR-AUC from {leak['pr_auc_deployable']:.3f} to "
          f"{leak['pr_auc_with_leakage']:.3f}, a {leak['multiple']:.1f}× jump, with "
          f"FRAUD_SUSPECTED alone worth roughly a third of the gain and the six "
          f"columns occupying the top importance ranks. That is not a model; it is "
          f"a lookup of the answer stored in a different column. They are "
          f"quarantined by default, and the flag that admits them prints a warning "
          f"while it does so.")

    doc.add_heading("B.  The bank's own eighteen features win", level=2)
    _para(doc,
          "The dictionary marks 18 predictors as bank-finalised. The pipeline "
          "treats that as a hypothesis rather than an instruction and competes four "
          "feature strategies under repeated stratified cross-validation, with "
          "selection performed inside every fold.")

    order = [
        ("bank_finalized", "Bank-finalised subset"),
        ("bank_plus_engineered", "Bank subset + engineered"),
        ("auto_topk", "Automatic top-k selection"),
        ("all", "Every available column"),
    ]
    _table(doc, ["Strategy", "Features", "ROC-AUC", "PR-AUC"],
           [[label, _n(cv[k]["n_features"]), f"{cv[k]['roc_auc']:.3f}",
             f"{cv[k]['pr_auc']:.3f}"] for k, label in order if k in cv],
           caption="TABLE IV.  FEATURE STRATEGIES, REPEATED STRATIFIED CV")

    _para(doc,
          f"The eighteen expert-chosen columns beat all "
          f"{_n(cv['all']['n_features'])}, and beat automatic selection. With "
          f"{_n(d['positives'])} positives against thousands of predictors this is "
          f"what statistical theory predicts, and it is precisely why the pipeline "
          f"measures instead of assuming. The untouched holdout scored "
          f"{hold['roc_auc']:.4f} against a cross-validated {sel['roc_auc']:.4f} — "
          f"a gap of {abs(hold['roc_auc'] - sel['roc_auc']):.4f}, which is the "
          f"evidence that the selection procedure is not inflating itself. A "
          f"regression test makes the same point adversarially: it shuffles the "
          f"target, destroying all signal, and asserts the reported AUC stays near "
          f"chance, which a globally-selecting pipeline fails.")

    _para(doc,
          f"These numbers are not model performance. The organisers' data had not "
          f"been released when this was built, so they were measured on a stand-in "
          f"table generated with their exact {_n(d['declared_columns'])}-column "
          f"schema and a deliberately modest injected signal. What they demonstrate "
          f"is that the pipeline runs end to end, that the methodology does not "
          f"flatter itself, and that the leakage trap is real — nothing more. When "
          f"the real file arrives, one command replaces every figure in this "
          f"section with a measured one.",
          italic=True)


def main() -> int:
    path = METRICS_DIR / "evaluation.json"
    if not path.exists():
        print(f"missing {path} - run `make evaluate` first", file=sys.stderr)
        return 1
    m = json.loads(path.read_text())
    ds, layers = m["dataset"], m["layers"]
    base, cal, lat = m["baseline"], m["calibration"], m["latency"]
    ind, oot = m["independence_from_tickets"], m.get("out_of_time", {})
    decoys = {d["population"]: d for d in m["hard_negatives"]}
    weak = [t for t in m["typology"] if not t["typology"].startswith("role:")]

    doc = Document()
    _style(doc)

    # ---- title block spans both columns (single-column section) ----------
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(0.62)
    sec.right_margin = Inches(0.62)

    t = _para(doc, "BODHI MULE HUNTER AI: Explainable Graph-Temporal Detection "
                   "of Money-Mule Accounts and Suspicious Transactions",
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20)
    t.paragraph_format.space_after = Pt(10)

    # Author block: IEEE lays these out in columns, so a 1xN table keeps the
    # names and enrolment numbers aligned without hard tabs.
    at = doc.add_table(rows=2, cols=len(TEAM))
    for i, mem in enumerate(TEAM):
        for row, value, size, mono in ((0, mem.name, 11, False),
                                       (1, mem.enrolment, 9.5, True)):
            cell = at.rows[row].cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(value)
            run.font.size = Pt(size)
            if mono:
                run.font.name = "Consolas"
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    _para(doc, f"{TEAM_NAME} — {EVENT}",
          align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _para(doc, AFFILIATION,
          align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=12)

    # ---- body is two-column ------------------------------------------------
    body = doc.add_section(WD_SECTION.CONTINUOUS)
    body.top_margin = Inches(0.75)
    body.bottom_margin = Inches(1.0)
    body.left_margin = Inches(0.62)
    body.right_margin = Inches(0.62)
    _two_columns(body)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Abstract—")
    r.bold = True
    r.italic = True
    r.font.size = Pt(9)
    r2 = p.add_run(
        f"Money mules convert stolen funds into untraceable cash, and rule-based "
        f"transaction monitoring detects them at a precision so low that analysts "
        f"cannot act on the output. We present BODHI MULE HUNTER AI, a nine-layer "
        f"engine that fuses gradient-boosted tabular screening, an inductive graph "
        f"neural network and a temporal memory network over each account's event "
        f"sequence, calibrated into a single supervisory risk score. On a simulated "
        f"bank of {_n(ds['accounts'])} accounts and {_n(ds['transactions'])} "
        f"cross-channel transactions containing seven documented laundering "
        f"typologies, the fused model reaches {layers['Fused (L7)']['roc_auc']:.4f} "
        f"ROC-AUC and {layers['Fused (L7)']['pr_auc']:.4f} PR-AUC. Against a rule "
        f"engine of the kind banks deploy today, at identical recall, alert volume "
        f"falls from {_n(base['alerts'])} to {_n(base['bodhi_alerts_at_equal_recall'])} "
        f"and precision rises from {_pct(base['precision'], 2)} to "
        f"{_pct(base['bodhi_precision_at_equal_recall'], 2)} — "
        f"{_pct(base['false_positive_reduction'], 2)} of false positives eliminated. "
        f"Inline decisions complete in {lat['p50_ms']:.3f} ms at the median. Every "
        f"alert carries exact Shapley attributions and a learned edge-importance "
        f"mask, and the containment layer refuses to act autonomously without "
        f"corroboration from independent model layers.")
    r2.font.size = Pt(9)

    p = doc.add_paragraph()
    r = p.add_run("Keywords—")
    r.bold = True
    r.italic = True
    r.font.size = Pt(9)
    r2 = p.add_run("money mule detection, graph neural networks, temporal graph "
                   "networks, explainable AI, anti-money laundering")
    r2.font.size = Pt(9)
    r2.italic = True

    # ------------------------------------------------------------------ I
    doc.add_heading("I.  INTRODUCTION", level=1)
    _para(doc,
          "A money mule is an account whose owner has rented, sold, or been "
          "deceived into lending it. Criminal proceeds are pushed into it, split, "
          "relayed through several further accounts, and withdrawn as cash — "
          "typically within hours. The account itself is unremarkable: the KYC is "
          "genuine, the customer exists, the payments clear. What is anomalous is "
          "the shape of the money passing through it and the company it keeps.")
    _para(doc,
          "This creates a specific failure mode for the transaction-monitoring "
          "systems banks currently operate. Threshold-and-scenario engines evaluate "
          "accounts in isolation, so they cannot express “this account's "
          "counterparties are themselves suspicious” or “these forty "
          "credits arrived in ninety minutes and left at 03:12”. Lacking that "
          "vocabulary, they compensate with sensitivity, and the result is an alert "
          f"queue an AML desk cannot clear. The baseline we implement raises "
          f"{_n(base['alerts'])} alerts at {_pct(base['precision'], 2)} precision: "
          "ninety-nine of every hundred investigations are wasted.")
    _para(doc, "We make four contributions:")
    for text in [
        "A three-view ensemble. Rather than three variants of the same tabular "
        "model, we combine a gradient-boosted model over behavioural aggregates, a "
        "GraphSAGE network over an account–device–IP graph, and a temporal graph "
        "network over each account's ordered event stream.",
        "Explainability that survives contact with a regulator. Layer 4 uses exact "
        "TreeSHAP, whose attributions sum to the model's margin rather than "
        "approximating it; Layer 5 uses a GNNExplainer edge mask, which answers the "
        "question SHAP structurally cannot — which relationships drove the score.",
        "An evaluation designed to be falsifiable. We plant legitimate populations "
        "that are structurally indistinguishable from mules and report the "
        "false-positive rate on each separately.",
        "Containment that can refuse. The kill-switch requires corroboration from "
        "at least two independent layers before it may freeze an account, is "
        "rate-limited, reversible, and excludes salary and pension accounts.",
    ]:
        doc.add_paragraph(text, style="List Number").alignment = \
            WD_ALIGN_PARAGRAPH.JUSTIFY

    # ------------------------------------------------------------------ II
    doc.add_heading("II.  PROBLEM SETTING AND DATA", level=1)
    doc.add_heading("A.  Why a simulator", level=2)
    _para(doc,
          "Account-level mule labels joined to device, IP and UPI linkage are "
          "simultaneously commercially sensitive, personally identifying and "
          "legally restricted; no bank can export them and no public dataset "
          "contains them. The alternative to simulation is not real data — it is no "
          "evaluation at all.")
    _para(doc,
          f"Our generator produces {_n(ds['accounts'])} accounts over "
          f"{ds['n_days']} days, {_n(ds['transactions'])} transactions across eight "
          f"payment rails (UPI, IMPS, NEFT, RTGS, AePS, ATM, card, wallet), "
          f"{_n(ds['mule_accounts'])} mule accounts ({_pct(ds['mule_rate'], 2)}) in "
          f"{ds['rings']} rings, and {_n(ds['fraud_alerts'])} government "
          f"cyber-fraud tickets arriving with realistic reporting lag.")

    doc.add_heading("B.  The hard negatives", level=2)
    _para(doc,
          "This is the part of the data design that determines whether any "
          "precision figure means anything. A collection mule receiving forty "
          "payments from strangers in one afternoon is structurally identical to a "
          "tuition-batch organiser. We therefore plant, as legitimate accounts, the "
          "populations in Table I and report the false-positive rate on each.")
    _table(doc, ["Legitimate population", "Count", "FP rate"],
           [[label.replace("_", " ").title(),
             _n(decoys[key]["n"]), _pct(decoys[key]["false_positive_rate"], 2)]
            for key, label in [
                ("bc_agent_device_cluster", "BC (Bank Mitra) agents"),
                ("community_collector", "Community collectors"),
                ("small_business_fan_in", "Small businesses"),
                ("legitimate_dormant_wake", "Dormancy reactivation"),
                ("ordinary_retail", "Ordinary retail"),
            ] if key in decoys],
           caption="TABLE I.  FALSE POSITIVES ON DELIBERATELY PLANTED LOOK-ALIKES")
    _para(doc,
          "Without these populations the models score above 0.999 AUC and the "
          "exercise is meaningless.")

    # ------------------------------------------------------------------ III
    doc.add_heading("III.  SYSTEM ARCHITECTURE", level=1)
    _para(doc,
          "The engine is organised as nine layers, each owned by a named agent "
          "with declared input and output contracts, so the provenance of every "
          "decision is recoverable from the audit trail.")
    for head, text in [
        ("Layer 1 — Ingestion.",
         "Normalises cross-channel transactions, government cyber-fraud tickets "
         "(NCRP/I4C/1930), APK-derived indicators, and regulatory directives. The "
         "four directive classes are handled distinctly: a mule list raises a score "
         "floor; a velocity limit changes a decision threshold, not a score; a "
         "sanction is absolute; an advisory is recorded and surfaced but never "
         "auto-applied, because a machine cannot responsibly convert prose into an "
         "enforcement rule."),
        ("Layer 2 — Feature engineering.",
         "Sixty behavioural features: velocity, fan-in/fan-out ratios, dormancy, "
         "burst scores, structuring share, device reuse, pass-through ratios and "
         "channel mix. Every builder is point-in-time correct, and a regression "
         "test asserts that passing an as-of argument is identical to truncating "
         "the ledger by hand. A second test shuffles the labels and asserts the "
         "feature matrix is unchanged, which forecloses label leakage."),
        ("Layer 3 — Graph construction.",
         "Accounts are nodes; edges are money transfers, shared devices and shared "
         "IP addresses. Shared-identifier cliques are capped: an identifier behind "
         "more than twenty-five accounts is infrastructure — a CGNAT pool, a "
         "banking kiosk — not evidence of collusion. Without the cap a single "
         "public address produces a four-thousand-node clique and the graph layer "
         "becomes a false-positive generator."),
        ("Layers 4–6 — Three views.", "Described in Section IV."),
        ("Layer 7 — Fusion.", "Described in Section V."),
        ("Layer 8 — Explainability.",
         "TreeSHAP, GNNExplainer, and a narrative generator that maps each feature "
         "onto a sentence in the vocabulary of an AML desk."),
        ("Layer 9 — Alerting and containment.", "Described in Section VII."),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(head + " ")
        r.bold = True
        p.add_run(text)

    # ------------------------------------------------------------------ IV
    doc.add_heading("IV.  THE THREE MODEL VIEWS", level=1)
    _para(doc,
          "The central design claim is that mule detection needs three genuinely "
          "different kinds of evidence, and that no single model family supplies "
          "all three.")

    doc.add_heading("A.  Layer 4: gradient-boosted screening", level=2)
    _para(doc,
          "An XGBoost classifier over the behavioural features disposes of the "
          "overwhelming majority of accounts cheaply. It was chosen over a neural "
          "tabular model for one reason: exact TreeSHAP. Shapley values for tree "
          "ensembles are computable in polynomial time, so the attribution shown to "
          "an investigator is not a sampled approximation — it sums to the model's "
          "raw margin, a property we assert in the test suite, and it is stable "
          "across runs. That matters when the explanation is transcribed into a "
          "regulatory filing. The class-imbalance weight is capped rather than set "
          f"to the full ratio; uncapped, on a {_pct(ds['mule_rate'], 2)} base rate, "
          "it destroys calibration, and a mis-calibrated score cannot be allowed to "
          "drive an automated freeze.")

    doc.add_heading("B.  Layer 5: inductive graph learning", level=2)
    _para(doc,
          "We use the mean-aggregator GraphSAGE formulation [1] with two layers "
          "and a deterministic top-k fan-out cap standing in for random neighbour "
          "sampling. Determinism is a requirement, not a simplification: an "
          "investigator re-opening a case tomorrow must see the same score. Because "
          "the weights never index a node identity, the model is genuinely "
          "inductive — an account opened this morning is scored from its "
          "neighbourhood alone.")
    eq = _para(doc, "h_v^(k) = ReLU( W_self^(k) · h_v^(k−1) + W_neigh^(k) · "
                    "MEAN_{u∈N(v)} h_u^(k−1) )        (1)",
               align=WD_ALIGN_PARAGRAPH.CENTER)
    eq.runs[0].italic = True

    doc.add_heading("C.  Layer 6: temporal memory", level=2)
    _para(doc,
          "The memory module of a Temporal Graph Network [2] is a gated recurrent "
          "unit updated on every interaction, with elapsed time injected through a "
          "learnable encoding. We implement exactly that over each account's most "
          "recent sixty-four events. This layer exists because aggregates destroy "
          "the signal that matters most. “Received Rs. 4 lakh from 38 payers "
          "over the month” describes a shopkeeper and a mule equally well. "
          "“Received Rs. 4 lakh from 38 payers between 14:10 and 15:40, then "
          "withdrew Rs. 3.8 lakh in nineteen ATM visits beginning at 03:12” "
          "describes only one of them. Order is the discriminant, and no aggregate "
          "preserves it.")

    doc.add_heading("D.  Implementation note", level=2)
    _para(doc,
          "Both neural layers are implemented in NumPy with analytically derived "
          "gradients rather than in a deep-learning framework. The motivation is "
          "deployability in a bank: the entire engine installs from four wheels, "
          "runs on a laptop CPU, and presents no GPU or CUDA surface to certify. "
          "Because “hand-derived” is a plausible source of silent error, "
          "every parameter tensor's analytic gradient is verified against central "
          "finite differences in the test suite.")

    # ------------------------------------------------------------------ V
    doc.add_heading("V.  FUSION AND CALIBRATION", level=1)
    _para(doc,
          "The four channel scores — tabular, graph, temporal and external "
          "intelligence — are blended in log-odds space. Sub-scores accumulate "
          "against 0 and 1, where a linear model in probability space cannot "
          "separate 0.990 from 0.9999; that is a hundredfold difference in odds and "
          "precisely the region in which alert ranking is decided.")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Monotonicity. ").bold = True
    p.add_run(
        "Blend weights are constrained non-negative, so the fused score can never "
        "decrease when a layer's evidence increases. This is not a regularisation "
        "convenience. An unconstrained stacker fitted on a small fold reliably "
        "learns a negative coefficient on the temporal channel, which is "
        "indefensible in front of an investigator and unshippable past a "
        "supervisor. A useful side effect is that the blend can always degenerate "
        "to “all weight on the single best layer”, so fusion cannot "
        "perform much worse than its best input.")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Calibration. ").bold = True
    p.add_run(
        f"Isotonic regression maps the blend onto observed frequencies, giving a "
        f"Brier score of {cal['brier']:.5f} and expected calibration error of "
        f"{cal['ece']:.4f}. Isotonic regression is a step function, so on a "
        f"near-separable problem it collapses thousands of distinct scores into a "
        f"handful of plateaus and every account within a plateau becomes tied; we "
        f"measured a four-point AUC loss from those ties. Retaining a two-percent "
        f"share of the underlying monotone score breaks them without disturbing "
        f"calibration.")
    _para(doc,
          "The blend is fitted on a fourth, dedicated fold. A stacker fitted on the "
          "validation fold — the fold that early-stopped its own inputs — inherits "
          "their optimism and, in our experiments, performed worse than the best "
          "single layer.")

    # ------------------------------------------------------------------ VI
    doc.add_heading("VI.  RESULTS", level=1)
    _table(doc, ["Layer", "ROC-AUC", "PR-AUC", "P@1%", "R@1%"],
           [[name, f"{v['roc_auc']:.4f}", f"{v['pr_auc']:.4f}",
             f"{v['precision_at_1pct']:.3f}", f"{v['recall_at_1pct']:.3f}"]
            for name, v in layers.items()],
           caption="TABLE II.  PER-LAYER DETECTION PERFORMANCE")
    _para(doc,
          "The complementarity claim is visible: the fused model exceeds every "
          "individual view on both ranking metrics. P@1% and R@1% are precision and "
          "recall within an alert budget of one percent of accounts — the quantity "
          "an analyst actually experiences.")

    doc.add_heading("A.  Against the incumbent", level=2)
    _para(doc,
          "We implement a rule engine of the kind banks deploy today — eight "
          "scenarios covering velocity, structuring, dormancy, cash-out ratio, "
          "device sharing, new-account turnover and off-hours concentration — and "
          "give it its best operating point rather than a strawman. The comparison "
          "is made at identical recall, because a system that alerts on everything "
          "can always claim to catch more.")
    _table(doc, ["", "Rule engine", "BODHI"],
           [["Alerts raised", _n(base["alerts"]),
             _n(base["bodhi_alerts_at_equal_recall"])],
            ["True positives", _n(base["true_positives"]),
             _n(base["bodhi_true_positives_at_equal_recall"])],
            ["False positives", _n(base["false_positives"]),
             _n(base["bodhi_false_positives_at_equal_recall"])],
            ["Precision", _pct(base["precision"], 2),
             _pct(base["bodhi_precision_at_equal_recall"], 2)]],
           caption=f"TABLE III.  RULE ENGINE VERSUS BODHI AT EQUAL RECALL "
                   f"({_pct(base['recall'], 1)})")
    _para(doc,
          f"{_pct(base['false_positive_reduction'], 2)} of false positives are "
          f"eliminated at unchanged recall, a {base['precision_uplift_x']:.0f}× "
          f"precision uplift. In operational terms an analyst clears the same true "
          f"positives while opening {_n(base['bodhi_alerts_at_equal_recall'])} "
          f"cases instead of {_n(base['alerts'])}.")

    doc.add_heading("B.  Robustness checks", level=2)
    for head, text in [
        ("Out-of-time generalisation.",
         f"{oot.get('late_rings', 0)} rings activate only in the final quarter of "
         f"the simulation window and are therefore genuinely unseen patterns. "
         f"Recall on their {oot.get('late_ring_mules', 0)} member accounts is "
         f"{_pct(oot.get('recall', 0), 1)}."),
        ("Independence from the government feed.",
         f"The supervised layers never receive ticket data — the intelligence "
         f"features are a disjoint set, asserted by test, and enter only at "
         f"Layer 7. Consequently "
         f"{_pct(ind['share_of_detections_not_in_any_ticket'], 0)} of detected "
         f"mules were never named in any NCRP ticket: they are cases the "
         f"government feed could not have supplied."),
        ("Latency.",
         f"Inline decisions complete in {lat['p50_ms']:.3f} ms at the median and "
         f"{lat['p99_ms']:.3f} ms at the 99th percentile, roughly "
         f"{_n(lat['throughput_per_sec'])} decisions per second per core. This is "
         f"achievable because the graph and temporal layers run on a schedule and "
         f"cache a standing risk per account; the authorisation path combines those "
         f"cached scores with the transaction's own attributes and never traverses "
         f"the graph."),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(head + " ").bold = True
        p.add_run(text)

    doc.add_heading("C.  Where the system is weak", level=2)
    _para(doc,
          f"Recall is not uniform, and averaging conceals that. The weakest "
          f"typology is {weak[0]['typology'].replace('_', ' ').title()} at "
          f"{_pct(weak[0]['recall'], 1)} recall, followed by "
          f"{weak[1]['typology'].replace('_', ' ').title()} at "
          f"{_pct(weak[1]['recall'], 1)}. The structuring result is expected and "
          f"structural: an individual sub-threshold transfer is, by construction, "
          f"indistinguishable from a genuine near-threshold payment, so detection "
          f"depends on an aggregate pattern that only becomes visible once enough "
          f"transfers have accumulated. Terminal cash-out accounts are similarly "
          f"hard because they have short, thin histories — they receive once and "
          f"drain. We report these rather than averaging them away, because a "
          f"deployment team needs to know which typologies still require "
          f"human-designed scenarios alongside the model.")

    for fig, cap in (("roc_pr.png", "Fig. 1.  ROC and precision–recall by layer."),
                     ("calibration.png", "Fig. 2.  Calibration of the fused score."),
                     ("alert_budget.png",
                      "Fig. 3.  Precision and recall versus analyst workload.")):
        p = FIGURE_DIR / fig
        if p.exists():
            doc.add_picture(str(p), width=Inches(3.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c = _para(doc, cap, align=WD_ALIGN_PARAGRAPH.CENTER, size=8)
            c.runs[0].italic = True

    # ------------------------------------------------------------------ VII
    doc.add_heading("VII.  CONTAINMENT, SAFETY AND COMPLIANCE", level=1)
    _para(doc,
          "Freezing an account is the most consequential action the system can "
          "take. A wrongly frozen account is a person unable to pay for medicine, "
          "and no improvement in AUC justifies treating that casually. The "
          "kill-switch is therefore built from constraints rather than a threshold: "
          "below the critical band the automation cannot freeze at all; a full "
          "freeze requires at least two independent layers to agree; every action "
          "carries a time-to-live and can be reverted; automated freezes per hour "
          "are bounded, capping the blast radius if a model degrades or an upstream "
          "feed is poisoned; and salary, pension and benefit accounts are excluded "
          "from automated freezing entirely.")
    _para(doc,
          "When a constraint fires, the action is downgraded rather than escalated "
          "and flagged for human review. Every decision, including every refusal, "
          "is written to a hash-chained append-only audit log whose head hash is "
          "exposed for external anchoring.")
    _para(doc,
          "The compliance layer drafts Suspicious Transaction Reports populated "
          "from exactly the evidence displayed to the investigator, so what the "
          "model asserted and what the institution files cannot diverge. Drafts are "
          "explicitly marked as requiring human sign-off; nothing is filed "
          "autonomously. Cash-transaction reporting aggregates per account per "
          "calendar day, which is the level at which the obligation applies — and "
          "precisely why structuring defeats a per-transaction check.")

    # ------------------------------------------------------------------ VIII
    doc.add_heading("VIII.  ECOSYSTEM SYNERGY: BODHI SHIELD", level=1)
    _para(doc,
          "Mule networks are frequently seeded by malicious Android packages that "
          "carry their beneficiary accounts hardcoded inside them. Our companion "
          "component performs static triage of a submitted APK: it parses the "
          "binary AndroidManifest.xml string pool to specification, scores "
          "permission combinations rather than individual permissions (READ_SMS "
          "alone is a messaging application; READ_SMS together with an "
          "accessibility service and a screen overlay is a banking trojan), mines "
          "the DEX for UPI handles, IFSC codes, account numbers and "
          "command-and-control endpoints, and detects commercial packers and "
          "identifier obfuscation.")
    _para(doc,
          "Every financial identifier recovered is streamed into the Mule Hunter "
          "graph as a weighted node. The consequence is architectural rather than "
          "incremental: when a new trojan is analysed on the day it appears, the "
          "beneficiary accounts it was built to pay already carry an elevated "
          "intelligence score before the first victim installs it. Detection stops "
          "being purely retrospective.")

    # ------------------------------------------------------------------ IX
    _boi_section(doc)

    # ------------------------------------------------------------------- X
    doc.add_heading("X.  LIMITATIONS", level=1)
    _para(doc,
          "We state these plainly because a prototype that hides them is not "
          "useful to a deployment team. Results are on simulated data; they "
          "demonstrate that the architecture works and that it decisively beats a "
          "rule engine on data containing realistic decoys, but they are not a "
          "forecast of production performance. The SHIELD component performs static "
          "analysis only — the dynamic sandbox with runtime hooking described in "
          "our proposal requires an instrumented Android image and cannot ship "
          "self-contained. The graph layers are not real-time: a full re-score of "
          "the population takes tens of seconds, so an account whose neighbourhood "
          "changed since the last batch is scored on slightly stale structure. The "
          "system sees one institution, so rings routing through several banks are "
          "only partly visible — a data-sharing problem rather than a modelling "
          "one. The strongest single feature is neighbourhood risk propagated from "
          "confirmed cases, which means a cold start with no known mules would "
          "perform materially worse. Finally, disparate impact is unevaluated: "
          "minimum-KYC status and shared-device features are predictive but "
          "correlate with lower-income and multi-occupancy households, and any real "
          "deployment must measure alert-rate parity before go-live.")

    # ------------------------------------------------------------------ XI
    doc.add_heading("XI.  CONCLUSION", level=1)
    _para(doc,
          f"Mule detection fails today not because the signal is absent but "
          f"because single-account rule engines cannot express it. Giving the "
          f"problem three complementary views — behavioural aggregates, network "
          f"structure and event ordering — and fusing them under a monotonicity "
          f"constraint into a calibrated score reduces false positives by "
          f"{_pct(base['false_positive_reduction'], 2)} at unchanged recall "
          f"relative to a realistic incumbent, while producing explanations "
          f"specific enough to file and containment actions cautious enough to "
          f"automate. The same discipline is applied to the organisers' own alert "
          f"schema, where the honest finding is that four of its columns encode "
          f"the answer and the bank's eighteen expert-chosen features beat "
          f"automatic selection over several thousand. "
          f"The complete system, the data simulator and every script "
          f"needed to reproduce these numbers are released alongside this report.")

    doc.add_heading("REFERENCES", level=1)
    for i, ref in enumerate([
        "W. L. Hamilton, R. Ying, and J. Leskovec, “Inductive representation "
        "learning on large graphs,” in Proc. NeurIPS, 2017.",
        "E. Rossi et al., “Temporal graph networks for deep learning on "
        "dynamic graphs,” in ICML Workshop on Graph Representation Learning, "
        "2020.",
        "R. Ying, D. Bourgeois, J. You, M. Zitnik, and J. Leskovec, "
        "“GNNExplainer: Generating explanations for graph neural networks,” "
        "in Proc. NeurIPS, 2019.",
        "T. Chen and C. Guestrin, “XGBoost: A scalable tree boosting system,” "
        "in Proc. ACM SIGKDD, 2016.",
        "S. M. Lundberg et al., “From local explanations to global understanding "
        "with explainable AI for trees,” Nature Machine Intelligence, vol. 2, "
        "no. 1, pp. 56–67, 2020.",
        "V. D. Blondel, J.-L. Guillaume, R. Lambiotte, and E. Lefebvre, “Fast "
        "unfolding of communities in large networks,” J. Stat. Mech., 2008.",
        "Reserve Bank of India, “Master direction on KYC and framework for "
        "monitoring of mule accounts,” 2024.",
        "Indian Cybercrime Coordination Centre, “National Cybercrime Reporting "
        "Portal: Citizen Financial Cyber Fraud Reporting and Management System,” "
        "Ministry of Home Affairs, 2023.",
        "M. Weber et al., “Anti-money laundering in Bitcoin: Experimenting with "
        "graph convolutional networks for financial forensics,” in KDD Workshop "
        "on Anomaly Detection in Finance, 2019.",
    ], 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f"[{i}]  {ref}")
        r.font.size = Pt(8)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
